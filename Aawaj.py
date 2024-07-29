# Import necessary libraries and modules
from tkinter import *
from tkinter import filedialog, messagebox
from tkinter import ttk 
import speech_recognition as sr
from pydub import AudioSegment
from docx import Document
import subprocess
import moviepy.editor
import shutil
import os
import threading 
import tempfile  
from tempfile import NamedTemporaryFile
from PIL import Image, ImageTk

# Function to convert video to audio
def convert_video_to_audio():
    global temp_audio_file  # Declare the global variable
    vid = filedialog.askopenfilename(filetypes=[("Video files", "*.mp4 *.avi *.mkv")])
    if vid:
        video = moviepy.editor.VideoFileClip(vid)
        aud = video.audio
        temp_audio_file = tempfile.NamedTemporaryFile(suffix=".mp3", delete=False).name  # Create a temporary file
        t = threading.Thread(target=convert_audio_thread, args=(aud, temp_audio_file))
        t.start()

# Threaded function to convert audio
def convert_audio_thread(aud, save_path):
    try:
        progress_bar.start()  # Start the progress bar
        aud.write_audiofile(save_path, codec="mp3", logger=None)
        progress_bar.stop()  # Stop the progress bar
        progress_bar["mode"] = "determinate"  # Switch to determinate mode
        progress_bar["value"] = 100  # Set the progress bar to 100%
        open_choice = messagebox.askyesno("Listen to MP3", "Do you want to listen to the MP3 file?")
        if open_choice:
            play_mp3_with_media_player(save_path)
        like_choice = messagebox.askyesno("Save the Song", "Do you want to save the mp3 file?")
        if like_choice:
            save_mp3_locally(save_path)
    except Exception as e:
        progress_bar.stop()  # Stop the progress bar on error
        messagebox.showerror("Error", "Failed to convert video to audio:\n" + str(e))

# Function to save MP3 locally
def save_mp3_locally(mp3_path):
    save_path = filedialog.asksaveasfilename(defaultextension=".mp3", filetypes=[("MP3 files", "*.mp3")], initialfile=os.path.basename(mp3_path))
    if save_path:
        shutil.copy(mp3_path, save_path)
        messagebox.showinfo("Save Complete", "MP3 file saved as " + save_path)
        try:
            os.remove(mp3_path)  # Attempt to remove the temporary file
        except Exception as e:
            messagebox.showerror("Error", "Failed to delete temporary MP3 file:\n" + str(e))
# Function to play MP3 with media player
def play_mp3_with_media_player(mp3_path):
    try:
        wmplayer_path = "C:\\Program Files (x86)\\Windows Media Player\\wmplayer.exe"  # Adjust the path as needed
        subprocess.Popen([wmplayer_path, mp3_path], shell=True)
    except Exception as e:
        messagebox.showerror("Error", "Failed to open MP3 with Windows Media Player:\n" + str(e))

# Function to convert WAV to text

# Function to convert audio to text
def convert_mp3_to_text():
    audio_file = filedialog.askopenfilename(filetypes=[("Audio files", "*.mp3")])
    if audio_file:
        wav_audio = convert_mp3_to_wav(audio_file)

        def recognize_audio():
            try:
                r = sr.Recognizer()
                with sr.AudioFile(wav_audio) as source:
                    r.adjust_for_ambient_noise(source)
                    audio = r.listen(source)
                    recognized_text = r.recognize_google(audio)

                    temp_docx_file = NamedTemporaryFile(suffix=".docx", delete=False)

                    save_text_as_word_document(recognized_text, temp_docx_file.name)

                    progress_bar_audio_mp3.stop()
                    progress_bar_audio_mp3["value"] = 100
                    open_choice = messagebox.askyesno("Open Word Document", "Do you want to open the Word document?")
                    if open_choice:
                        open_word_document(temp_docx_file.name)
                        like_choice = messagebox.askyesno("Save the Document", "Do you want to save the Word document?")
                        if like_choice:
                            save_docx_locally(temp_docx_file.name)
            except Exception as e:
                progress_bar_audio_mp3.stop()
                progress_bar_audio_mp3["value"] = 100
                status_label.config(text="Conversion failed.")
                messagebox.showerror("Error", "Failed to convert audio to text:\n" + str(e))

        progress_bar_audio_mp3["value"] = 0
        progress_bar_audio_mp3.start()

        t = threading.Thread(target=recognize_audio)
        t.start()

# Function to convert MP3 to WAV using pydub
def convert_mp3_to_wav(mp3_path):
    audio = AudioSegment.from_mp3(mp3_path)
    wav_audio = NamedTemporaryFile(suffix=".wav", delete=False).name
    audio.export(wav_audio, format="wav")
    return wav_audio

# Function to save Word document locally
def save_docx_locally(docx_path):
    save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word documents", "*.docx")], initialfile=os.path.basename(docx_path))
    if save_path:
        shutil.copy(docx_path, save_path)
        messagebox.showinfo("Save Complete", "Word document saved as " + save_path)
        try:
            os.remove(docx_path)
        except Exception as e:
             messagebox.showerror("Error", "Failed to delete temporary docx file:\n" + str(e))    
# Function to save recognized text as Word document
def save_text_as_word_document(text, docx_filename):
    doc = Document()
    doc.add_heading("Recognized Text", 0)
    doc.add_paragraph(text)
    doc.save(docx_filename)

# Function to open Word document
def open_word_document(docx_filename):
    try:
        subprocess.Popen(["start", "winword", docx_filename], shell=True)
    except Exception as e:
        messagebox.showerror("Error", "Failed to open the Word document:\n" + str(e))
      

# Initialize the Tkinter window
root = Tk()
root.title("Aawaj")
root.geometry("700x450")
root.minsize(700, 450)
root.maxsize(700, 450)
frame = Frame(root)
frame.pack(side=BOTTOM, fill=BOTH, expand=True)

# Load and display background image
bg_image = Image.open("C:\\Users\\vbs\\bg.png")
bg_photo = ImageTk.PhotoImage(bg_image)
canvas = Canvas(root, width=700, height=450)
canvas.pack()
canvas.create_image(-75, 0, anchor=NW, image=bg_photo)

# Display left image and text
left_image = Image.open("C:\\Users\\vbs\\v-a.png")
left_photo = ImageTk.PhotoImage(left_image)
canvas.create_image(50, 200, anchor=NW, image=left_photo)
canvas.create_text(347, 350, text="Now convert your videos to mp3 files and also mp3 files to text with ease", font=("Lucida Calligraphy", 12, "bold"), fill="black")

# Display right image and title
r_image = Image.open("C:\\Users\\vbs\\a-t.png")
r_photo = ImageTk.PhotoImage(r_image)
canvas.create_image(520, 210, anchor=NW, image=r_photo)
title_label = Label(canvas, text="Aawaj", font=("Arial", 36, "bold"),bg="white",padx=5,pady=2 )
title_label.place(relx=0.5, rely=0.1, anchor="n")  # Position at the top center
canvas.create_text(150, 150, text="Video to audio converter", font=("Arial", 18, "bold"), fill="black")
canvas.create_text(560, 150, text="Audio to text converter", font=("Arial", 18, "bold"), fill="black")

# Create frames for buttons and progress bars
video_frame = Frame(frame)
video_frame.pack(side=LEFT, padx=5, pady=5)

audio_frame = Frame(frame)
audio_frame.pack(side=RIGHT, padx=5, pady=5)

# Create progress bars
progress_bar_audio_mp3 = ttk.Progressbar(audio_frame, length=200, mode="determinate")
progress_bar_audio_mp3.pack(side=TOP, padx=5, pady=5)

progress_bar = ttk.Progressbar(video_frame, length=200, mode="determinate")
progress_bar.pack(side=TOP, padx=5, pady=5)

# Create buttons for video and audio conversion
convert_video_button = Button(video_frame, text="Convert Video to Audio", command=convert_video_to_audio)
convert_video_button.pack(side=TOP, anchor="s")

convert_audio_button = Button(audio_frame, text="Convert MP3 to Text", command=convert_mp3_to_text)
convert_audio_button.pack(side=TOP, anchor="s")

# Label to display optional status message
status_label = Label(frame, text="", anchor="w")
status_label.pack(side=TOP, fill=X, padx=5, pady=5)

# Start the Tkinter main loop
root.mainloop()
